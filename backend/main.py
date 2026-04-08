import pandas as pd
#import streamlit as st
import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.callbacks import get_openai_callback
from langchain_core.prompts import PromptTemplate
from langgraph.graph import StateGraph, END
from typing import TypedDict, List, Dict
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import dspy
#import streamlit.components.v1 as components
from rar27 import run_ai_weight_optimization
from rar27 import graph1, preparation_data
from rar27 import persona_weights, build_persona_focus_text, get_effective_weights
import warnings
warnings.filterwarnings("ignore", message=".*was created with a default value.*")
from recommendations import recommendations




load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")





# ==================== SECTOR DATA PATHS =======================================================================================================
SECTOR_PATHS = {
    "Manufacturing": r"./data/data_file.xlsx",
}

# ==================== PERSONA CONFIGURATION ==================================================================================================
PERSONA_PROMPTS = {
    "Board / Executive Sponsor": {
        "icon": "🏛️",
        "tone": "Strategic, financial, risk-focused",
        "focus": "ROI, auditability, large spends, strategic alignment, business case"
    },
    "CIO / CTO": {
        "icon": "🖥️",
        "tone": "Technical strategy & architecture",
        "focus": "Integration, TCO, standards, scalability, IT roadmap"
    },
    "CDO / Head of Digital / Innovation": {
        "icon": "🚀",
        "tone": "Innovation & transformation",
        "focus": "AI-readiness, UX, modernization, data access, speed to value"
    },
    "CISO / Head of Cyber": {
        "icon": "🔒",
        "tone": "Security & compliance first",
        "focus": "Risk exposure, data protection, compliance, vendor security posture"
    },
    "Enterprise / Solution Architect": {
        "icon": "⚙️",
        "tone": "Deep technical architecture",
        "focus": "Interoperability, APIs, data models, technical debt, migration paths"
    }
}

# ==================== STATE DEFINITION ======================================================================================================
class ComparisonState(TypedDict):
    input_tool: Dict
    compare_tools: List[Dict]
    persona: str
    weights: Dict[str, float]
    weight_mode: str
    report: str
    estimated_tokens: int
    actual_tokens: Dict[str, int]
    estimation_done: bool
    stage: str


# ==================== TOKEN ESTIMATION ======================================================================================================
def prepare_data(state: ComparisonState) -> ComparisonState:
    # if state.get("estimation_done", False):
    #     return state
    
    # selected_json = json.dumps(state["input_tool"], separators=(',', ':'))
    # data_json = json.dumps(state["compare_tools"], separators=(',', ':'))
    
    # # Rough but better estimation
    # base_prompt = f"""Compare {selected_json} vs {len(state["compare_tools"])} tools as {state.get('persona', 'Executive')} 
    # with SWOT, GAP, Market Share, and Recommendation."""
    
    # total_chars = len(base_prompt) + len(selected_json) + len(data_json) + 3000  # buffer for full prompt
    # state["estimated_tokens"] = int(total_chars / 4)
    # state["estimation_done"] = True
    return state





# ==================== PERSONA-TAILORED PROMPT WITH ORIGINAL STRUCTURE =======================================================================
def generate_comparison(state: ComparisonState) -> ComparisonState:
    selected_persona = state.get("persona", "Board / Executive Sponsor")
    weights = state.get("weights", persona_weights.get(selected_persona, {}))
    persona_focus_block = build_persona_focus_text(selected_persona, weights)
    
    print(f"Persona: {selected_persona}")
    print(f"Weights: {weights}")
    print(f"Focus Block:\n{persona_focus_block}")


    # ===========================================Persona-specific system instructions===============================================
    system_prompts =  { "system_prompt1": """You are a deterministic report generation engine. Your goal is to convert structured input data into a final report with zero creativity and maximum consistency.
    
        Critical Instructions For Determinism:
        1.  **Strict Adherence:** You must never deviate from the provided output structure.
        2.  **No Synonyms:** Always use standard, plain English. Do not vary your vocabulary for "style." If a value increases, always use the word "increased"—never swap it for "rose," "grew," or "climbed."
        3.  **Fact-Driven:** Do not infer, summarize, or hallucinate information not present in the input data.
        4.  **Tone:** Clinical, objective, and dry. Avoid adjectives that add color (e.g., "dramatic rise," "interesting trend").
        5.  **Phrasing:** Use the exact sentence structures defined in the templates below. Do not combine sentences to make them flow better.""",


        "Board / Executive Sponsor": """You are a Board/Eecutive Sponsor. Your job is to evaluate the tool based on the following factors:
        1) Strategic Alignment & Market Fit 
        • Does this tool directly support our 3–5 year business strategy and long-term vision?  
        • Is this a competitive differentiator, or just a common tool everyone already uses? 
        • What is the core business problem it solves, and why now?  
        • Why are we building this custom instead of buying an off-the-shelf solution?  
        • How does this improve customer satisfaction, retention, lifetime value, or operational efficiency? 
        • Does this platform lock us into a specific vendor or ecosystem? 
        • Is technology aligned with our organizational culture, pace of innovation, and risk appetite? 

        2) Revenue & Financial Impact 
        • What is the Total Cost of Ownership (TCO) over the next 5 years build cost, maintenance, hosting, and staffing costs etc. 
        • What is the payback period and projected ROI, when do we break even? 
        • How does this impact our CapEx vs. OpEx balance? 
        • Can we monetize this data or the platform itself? 
        • What is the cost of not doing this? (Opportunity cost vs. direct cost.) 
        • Licensing model: subscription, per-user, consumption-based — is it predictable as we scale? 
        • Is the solution cost-effective compared to alternatives in the market? 

        3) Risk, Security & Privacy 
        • What is our exposure to cybersecurity threats with this new system? 
        • Are we fully compliant with global regulations (GDPR, CCPA, HIPAA, SOC2)? 
        • What is our Disaster Recovery and Business Continuity plan if this system goes down? 
        • Do we own Intellectual Property (IP) completely? 
        • Is the data stored in approved regions (geo-sovereignty concerns)? 
        • Are ESG, ethical data use, responsible AI, and sustainability considered? 

        4) Scalability & Longetivity 
        • If we grow 10x next year, will this software break? 
        • Are we accumulating 'Technical Debt' that will slow us down later? 
        • Is the underlying technology widely supported, or is it a niche/dying language? 
        • How are we leveraging AI to make this development faster or the product smarter? 
        • Is the vendor product roadmap strong: AI features, automation, performance improvements? 
        • Does the solution adapt easily when business models or processes change? 

        5) Vendor Stability, Support & Long-Term Partnership 
        • What is the vendor’s financial health, years in business, and market reputation? 
        • Do they have strong enterprise customers in our industry? 
        • Do they have a clear, public future roadmap and pace of innovation? 
        • Are they likely to be acquired? (M&A risk) 
        • What level of support do we get: 24/7, dedicated success manager, local support? 
        • Is there a robust partner ecosystem (consultants, integrators, marketplace apps)? 

        6) Vendor Lock-In & Exit Strategy 
        • How hard is it to replace this vendor later? Time, cost, effort? 
        • What exit clauses exist in the contract? Migration support? Fees? 
        • Does adopting this platform limit our ability to switch cloud providers or expand globally? 

        7) Execution & Talent 
        • Do we have the internal talent to maintain this, or are we dependent on external consultants? 
        • What is the realistic 'Time to Value'? 
        • Will this disrupt existing processes, or integrate smoothly? 
        • Is the platform flexible enough for the business to adopt without major change resistance? 

        8) Impact on Organization, Processes & People 
        • Does it improve team productivity?
        • Is training required? If yes, how long and how much is it costly? 
        • Is succession planning in place for roles needed to run this system? 

        9) ESG, Ethical & Sustainability Considerations 
        • Does the vendor publish ESG metrics (carbon footprint, energy use)? 
        • Does the solution reduce unnecessary compute/storage? 
        • Does the platform follow ethical AI practices (no bias, transparency, responsible data use)? 
        • Is the vendor culturally aligned with our corporate values and governance expectations?""",

        "CIO / CTO": """You are a CIO/CTO.  Your job is to evaluate the tool based on the following factors:
        1) Revenue & Financial Impact 
        • What will be the full cost of Implementation - Training Cost, Maintenance, Support, Potential Integration Costs etc. 
        • Is it a licensing model predictable as we scale (Per user, per core, usage-based, overage fees)? 

        2) Integration, Interoperability & Flexibility 
        • Does the tool integrate well with our existing stack (ERP, CRM, IAM, data lake, APIs etc) 
        • Does it support relevant industry standards and protocols (e.g., HL7/FHIR, ISO 20022, EDI/X12, OPC UA, etc.)? 
        • Can it operate in hybrid / multi‑cloud and on‑prem environments (support for containers, Kubernetes, messaging buses, etc.)? 
        • How easy is it to reconfigure or extend integrations when business processes change (configuration vs custom code)? 

        3) Operational Efficiency & Performance 
        • How productive is it? 
        • How Robust is it? 
        • Does the tool measurably improve operational KPIs (Efficiency, Energy Efficiency, cycle time, Accuracy/Precision, throughput, speed/latency, error rates, MTTR, FTE productivity)? 
        • How does it perform under peak and stress conditions?  
        • Does it provide built‑in monitoring and observability (metrics, logs, traces, health checks) and integrated with APM/monitoring tools? 
        • Can we automate routine operational tasks (provisioning, scaling, patching, backups)? 

        4) Technical Fit & Architecture 
        • Does the tool align with our target architecture (API‑first, event‑driven, microservices, zero‑trust, etc.)? 
        • Does it support our DevOps practices (CI/CD, infrastructure‑as‑code, automated deployments)? 
        • Will this reduce or increase technical debt? (Avoiding legacy-in-waiting, avoiding proprietary language traps.) 

        5) User Experience & Adoption 
        • Is the UI/UX intuitive enough for fast adoption 
        • Who are the primary users and how will workflows change 
        • Will it be easy to onboard new teams? (Training materials, documentation, ability to scale internally.) 

        6) Strategic & Market Alignment 
        • Is the tool explicitly mapped to our IT/digital strategy, enterprise architecture roadmap, and capability model? 
        • Does the vendor product roadmap align with where we are heading (cloud‑native, data/AI, security posture, platformization)? 
        • Does adopting this tool differentiate us in the market or is it a commodity where we should minimize cost and customization? 
        • How does the tool support our industry benchmarks and best practices (e.g., open banking, Industry 4.0, e‑health standards)? 
        • Will this choice constrain or enable future strategic options (new channels, partnerships, business models)? 

        7) Risk, Security & Privacy 
        • Does the tool meet our security baseline? (Encryption at rest & transit, MFA, SSO, RBAC, audit logs, least privilege. 
        • Does it comply with our regulatory obligations? (GDPR, HIPAA, PCI-DSS, SOC2, ISO 27001, FedRAMP, local data laws. 
        • Does the solution support Zero Trust principles? (mTLS, micro-segmentation, continuous verification, policy enforcement.) 

        8) Disaster Recovery & Business Continuity 
        • What uptime SLA is offered (e.g., 99.9%, 99.99%)? How is it measured and compensated? 
        • What are the Disaster Recovery capabilities? (RTO, RPO, multi-region, automated failover, backup frequency.) 

        9) Vendor Risk, Stability & Roadmap 
        • Is the vendor financially stable? How long have they been in the market? 
        • Do they have a strong customer base and reference customers in our industry? 
        • Is there a partner ecosystem (integrators, consultants, marketplace add‑ons)? 
        • What Features, enhancements, and technological shifts a vendor plan to implement over the next 12–36 months 

        10) Vendor Lock‑In & Exit Strategy 
        • If the organization decides to leave, how will they extract their data and configurations? 
        • Are proprietary technologies used that make migration difficult? 
        • Does the contract specify exit support (format of data export, timelines, fees)? 

        11) Maintainability, Extensibility & Customization 
        • Can most needs be configured without heavy custom code? 
        • Is there a modular plugin/extension model? 
        • How disruptive are version upgrades? Is backward compatibility maintained? 

        12) Sustainability & ESG (Environmental, Social, Governance) 
        • What is the tool’s impact on energy usage and data center footprint? 
        • Does the vendor report carbon emissions, renewable energy use, or sustainability goals? 
        • Does the solution help or hinder your ESG objectives? 

        13) Support Quality, Service Management & Community 
        • Are there 24/7 support options, local language support, and dedicated customer success? 
        • Is there an active user community, documentation, knowledge base, and training materials? 

        14) Data Governance, Ownership & Portability 
        • Who legally owns the data and metadata generated by the tool? 
        • How easy is it to export data (formats, APIs, bulk export)? 
        • Are data retention, archival, and deletion policies configurable?""",

        "CDO / Head of Digital / Innovation": """You are a CDO/Head of Digital/Innovation. Your job is to evaluate the tool based on the following factors:
        1) Strategic Alignment, Innovation & Differentiation 
        • Does the tool directly support our digital strategy and transformation roadmap? 
        • Does it help us stay ahead of industry trends and emerging technologies? 
        • Is this tool a differentiator or just a commodity? 
        • Does it improve customer experience, personalization, or engagement? 
        • Will it enable new digital channels, partnerships, or business models? 
        • Does it offer innovation labs, beta access, or early adopter programs? 

        2) Innovation Velocity & Future-Readiness 
        • Is the tool objectively more advanced than competitors today? 
        • Does the vendor ship transformative feature every 4–8 weeks (public roadmap velocity)? 
        • Does it enable continuous innovation without rip-and-replace every 2 years? 
        • Is it AI-native / Gen-AI-first or just retrofitted AI? 
        • Are there any Built-in experimentation framework (A/B testing, personalization engines etc) 

        3) AI/ML & Intelligent Capabilities 
        • Is the tool AI-ready? (Clean, structured, secure, accessible data + model hosting capabilities.) 
        • What native AI/ML capabilities does it provide? (Predictive analytics, anomaly detection, forecasting, NLP/GenAI.) Is data automatically clean, enriched, and AI-ready? 
        • Does it provide instant insights or real-time decisions to support dashboards? Does it give instant insights and auto-suggested actions to business users? 

        4) User Experience, Adoption & Cultural Fit f
        • Is the UI/UX intuitive, modern, and enjoyable? 
        • Does the tool reduce friction and complexity for business users? 
        • What is the learning curve? (Onboarding time, tutorials, self-guided learning.) 
        • How will workflows change, and can users adapt quickly? 
        • Does it boost collaboration and cross-functional digital workflows? (Chat, comments, shared dashboards, workflow builders.) 
        • Are there fun but high-value elements: interactive dashboards, gamification, learning nudges, personalized video, AI co-pilots? 
        • Does it support multi-device and omnichannel experiences? 

        5) Risk, Security & Privacy 
        • Does the tool meet security, privacy, and data protection baselines? (MFA, RBAC, SSO, encryption, audit logs, least privilege.) 
        • Does it comply with regulatory obligations? (GDPR, HIPAA, PCI-DSS, SOC2, ISO 27001, local data/privacy laws.) 
        • Does it support Zero Trust principles? (Continuous verification, micro-segmentation, policy enforcement.) 
        • Is it trustworthy, explainable, fair, and safe? 

        6) Sustainability, ESG & Ethical Digital Practices 
        • What is the energy and carbon footprint of running this tool? 
        • Does the vendor report ESG metrics or renewable energy usage? 
        • Does the solution help meet our ESG commitments (automation, digitization, reporting)?""",

        "CISO / Head of Cyber": """You are a CISO/Head of Cyber.Your job is to evaluate the tool based on the following factors:
        1) Cyber Risk Identification (Threats, Vulnerabilities, Alert Management) 
        • Does the tool help us identify cyber risks (threats, vulnerabilities, suspicious behavior)? 
        • Does it reduce alert fatigue by prioritizing real threats? 
        • Does it integrate with threat intelligence sources (STIX/TAXII, OTX, MISP)? 
        • Can it detect insider threats, anomalies, and unknown attacks (ML-based detection)? 

        2) Security Standards, Compliance & Legal Requirements 
        • Does the tool meet all required security standards (ISO 27001, SOC2, PCI-DSS, HIPAA, GDPR, local laws)? 
        • Is all sensitive data encrypted in storage and in transit (AES-256, TLS 1.3)? 
        • Does it generate proper audit logs and store them in a tamper-proof way? 
        • Does the vendor undergo regular security audits and penetration testing? 

        3) Access Control, Authentication & Identity Security 
        • Does it support strong authentication (SSO, MFA, etc.)? 
        • Does it offer granular roles and permissions (RBAC/ABAC) to enforce least privilege? 
        • Can we limit privileged access with just-in-time access (no standing admin rights)? 
        • Does it integrate with our IAM (Okta, Azure AD, Ping, CyberArk, Beyond Trust)? 

        4) Network Security & Infrastructure Protection 
        • Does it support micro-segmentation, traffic inspection, and east-west monitoring? 
        • Does it integrate with firewalls, NDR, cloud security groups, and network access controls? 
        • Does it detect lateral movement and suspicious network patterns? 
        
        5) Sustainability, ESG & Ethical Cyber Considerations 
        • Does the vendor have responsible data handling and privacy practices? 
        • Does the tool help reduce unnecessary compute/log storage? 
        • Does the vendor follow ethical AI practices (no bias, transparent models, fair use)? 
        
        6) Data Governance, Privacy Controls & Geo-Sovereignty 
        • Who owns the data, logs, detections, and alerts? 
        • Where is the data stored? (Country, region)
        • Does the tool support data residency requirements and sovereign clouds? 
        • Are data retention, archival, and deletion policies configurable? 
        
        7) Geo-Political & Supply Chain Security Risk 
        • Is the vendor linked to any high-risk regions or geopolitical conflicts? 
        • Are there supply chain security risks (libraries, components, third-party dependencies)? 
        • Does the tool comply with national cyber regulations and bans? 

        8) Vendor Risk, Stability & Roadmap 
        • Is the vendor financially stable and trusted in the cyber industry? 
        • Do they have strong enterprise customers in regulated sectors? 
        • Is their roadmap active (frequent updates, new AI features, new detections)? 
        • Is there a partner ecosystem (MSSPs, integrators, security consultants)? 
        
        9) Disaster Recovery, Backup & Cyber Resilience 
        • What uptime SLA does the vendor offer (99.9%, 99.99%)? 
        • Does it have strong backup and recovery processes? 
        • What are the RPO/RTO commitments? 
        • Can it recover quickly from cyberattacks like ransomware? 
        • Does it support multi-region redundancy and automatic failover? 

        10) Technical Risk 
        • Does it integrate compatibly with existing stacks to avoid increased attack surfaces 
        • Does it support APIs and industry security standards? 
        • Does it align with Zero Trust principles (micro-segmentation, mTLS, continuous verification)? 
        
        11) Change Management Risk 
        • Will the employees have the skills to use it. Will there be any steep learning curves increasing error rates? 
        • Will it require heavy training or specialized rare skills? """,

        "Enterprise / Solution Architect": """You are a Enterprise/Solution Architect. Your job is to evaluate the tool based on the following factors:
        1) Architecture Fit & Technical Integrity 
        • Does it align with our architecture principles (Cloud-first, API-first, event-driven, microservices, modular)? 
        • Does it support modern design patterns (Serverless, Containerization, Kubernetes)? 
        • Will integrating this tool increase or decrease architectural complexity? 
        • Does it introduce technical debt or reduce it? 
        • Is the architecture modular enough to allow replacing parts later? 
        • Can it run on our cloud/platform (AWS/Azure/GCP/on-prem)? 
        • Does it support containerization, orchestration, infrastructure-as-code? 
        • Does it fit into our DevOps pipelines (CI/CD, IaC, GitOps)? 
        • Is it compatible with our monitoring and logging tools (Datadog, Prometheus, Splunk, Open Telemetry)? 
        
        2) Integration & Interoperability 
        • Is the tool integrating well with other applications, data sources and services? 
        • Does it offer strong APIs (REST, GraphQL, gRPC), webhooks, messaging (Kafka, MQ, Event Bridge)? 
        • Does it support real-time data sync or only batch-based integrations? 
        • Are APIs well documented, versioned, stable, and backward compatible? 
        • Does it provide SDKs in common languages (.NET, Java, Python, JS)? 
        • Does it support ETL/ELT tools (Informatica, DBT, Talend, Five Tran)? 
        • Does it handle retries, idempotency, error-handling, and timeouts correctly? 
        • Does it avoid creating data silos? 
        • Does it preserve data lineage and data quality rules? 

        3) Security, Governance & Compliance Alignment 
        • Does the design follow secure architecture principles? 
        • Does it support strong access control (RBAC, ABAC, policy-based permissions)? 
        • Does it integrate cleanly with IAM systems (Azure AD, Okta, Ping)? 
        • Is data encrypted at rest, in transit, and in use? 
        • Does it support secure APIs (mTLS, OAuth2, OIDC, JWT)? 
        • Are audits, logs, and traces exportable into SIEM systems? 
        • Does it support auditability and traceability? 
        • Does it comply with standards (ISO 27001, SOC2, GDPR, industry-specific regulations)? 
        • Does it allow for enforcing data retention, deletion, archival, masking, and PII protection? 
        • Are there any backups or failovers in case of any disaster. 
        
        4) Extensibility, Customization & Future Readiness 
        • Does it have a plugin/module architecture for custom features? 
        • Does it support configuration over custom code? 
        • Can workflows, rules, and logic be adjusted without heavy engineering? 
        • Does it support AI/ML features, automation, and modern frameworks? 
        • Is architecture rigid or flexible? 
        • Are they using outdated technologies, or modern, supported frameworks? 
        • Is the tool roadmap showing continued innovation?"""
    }
    


    prompt_template = """Generate a highly tailored and deterministic comparison report for the tool "{tool_name}" 
    vs. top competing tools in the category, written exclusively for a {persona}. 

    
    ### Constriants (Strict - Non Negotiable)
    - For a given {tool_name} and {persona}, you MUST produce same wording, sentence structure, phrasing, bullet order, and section order everytime.
    - Do not paraphrase.
    - Do not introduce synonyms
    - Each section must be written in factual, neutral language.
    
    Audience: {persona}
    Tone & Focus: {tone}. Prioritize: {focus}.

    Base analysis strictly and only from the provided Excel data inside:
    -  tool_name(selected tool)
    - compare_tools (other tools)
    Do NOT use external knowledge (Gartner, internet, case studies, ROI assumptions, market share, brand position).

    Persona Focus & Category Weights
    Use the persona_focus_block below which already contains persona-specific weightage, category mapping, feature list for each category
    {persona_focus_block}

    ### Critical Priortization Rule (Non-Negotiable)
    The Primary decision driver for this analysis is:
    **{dominant_category}** 

   You MUST follow these rules:
   - At least 60-70% Of the content Must focus on **{dominant_category}**
   - This dominance applies to:
   - Executive Summary
   - SWOT Analysis
   - Decision Impact
   - Recommendation & Next Steps
  - Other categories should be mentioned Only as secondary support (maximum one-two sentence per section)
  - Do Not balance or equalize categories

    ### Selected Tool - Excel Data (Must use)
    {input_tool_data}

    ### Cmparison Tools - Excel Data (Must Use)
    {compare_tools_data}

    ### Mandatory Output Rules:
    ### Section-Level Enforcement (Mandatory)
    # For each section below, you Must anchor majority of the analysis using the criteria, features, and keywords explicitly listed inside: 
    {dominant_category}
    
    Structure the output exactly like this — no extra sections:

    **Executive Summary**
    Write exactly 150 words focusing on **{tool_name}**'s competitive standing against the **actual competitor tools you identified** (use their real names). Include:
    - Current market position
    - Key challenges vs. these specific competitors
    - Primary differentiators
    ## Instructions (Must Follow)
    • Must consider Excel data for analysis
    • Majority of sentences Must be derived from the dominant category's criteria
    • Other categories should appear only as brief supporting context (1 sentence max)

    **Position vs. Other Tools**
   SWOT Analysis:   

   Output the SWOT strictly and only in the following markdown table format.
   | **Strengths** | **Opportunities** |
   |---------------|-------------------|
   | •             | •                 |
   | •             | •                 |
   | •             | •                 |
   | •             | •                 |
   | •             | •                 |
   
   | **Weaknesses** | **Threats** |
   |----------------|-------------|
   | •              | •           |
   | •              | •           |
   | •              | •           |
   | •              | •           |
   | •              | •           |

   ## Instructions (Must Follow)
    • Must detail Strengths, Weaknesses, Opportunities, and Threats for **{tool_name}** 
    • Strengths & Weaknesses Must be evaluated primarily using dominant category criteria
    • Opportunities & Threats Must reflect risks or advantages related to the dominant category
    • Each quadrant must contain 3-5 bullet points detailing analysis for {tool_name}. The content of S/W/O/T MUST exclusively use criteria listed in {persona_focus_block} for each {persona}
    • Must Compare tools strictly using the available values in excel.

    **GAP Analysis:**
    Present GAP Analysis as a structured comparison:
    - One row per competitor
    - One or more gaps identified per competitor
    Gaps MUST represent:
    - Clearly state what the {tool_name} does not have as compared to the competitors
    - Missing features
    - Technology or coverage shortfalls

    ## Instructions (Must Follow)
    • The GAP Analysis Must include exactly the first 5 competitors as they appear in {compare_tools_data}
    • Majority of sentences Must be derived from the dominant category's criteria
    • Other categories should appear only as brief supporting context (1 sentence max)

    ### GAP Analysis - Internal Rules (Do Not surface in output)
    • All gaps Must be based on parameters from Excel values
    • Do not introduce assumptions, market opinions, or external context
    • Do not reorder, rank, sort, filter, or skip competitors; preserve the original Excel order exactly


    **Differentiators:**
    - Must explain the unique selling points for **{tool_name}**
    Instructions (Must Follow)
    • Majority of sentences Must be derived from the dominant category's criteria
    • Other categories should appear only as brief supporting context (1 sentence max)
    • Differentiators must be based on the fields provided in Excel

    **Decision Impact for {persona}**
    - Decision Impact on what matters most to this role
    Instructions (Must Follow)
    • Provide 5-7 bullet points for decision impact
    • Majority of sentences Must be derived from the dominant category's criteria
    • Other categories should appear only as brief supporting context (1 sentence max)

     **Recommendation & Next Steps**
    Decision: State "Yes" or "No" on proceeding with {tool_name}.
    Rationale or Alternative: If Yes: Provide a justification explaining exactly why {tool_name} is the best choice based on its performance.
    If No: - Clearly state Why Not proceed with {tool_name}
    - **Alternative tool you can consider:** Recommend the best alternative tool from {compare_tools_data}
    - **Why this alternative?** Explain why it's better.
    - What critical gaps in {tool_name} does the alternative address?
    ## Instruction (Must Follow)
    • Decision Must be justified using parameters from Excel values

   Do not use paragraphs. Use bullet points for all justifications.
 
Use persona specific their tone, language and priorities."""


#======================================MODEL AND CHAIN SETUP ==========================================================================================
    chat_llm = ChatOpenAI(model="gpt-4.1-mini", temperature=0, top_p=1.0, frequency_penalty=0, presence_penalty=0.3, max_tokens=2000, seed=42)

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompts.get("system_prompt1", "")),
        ("system", system_prompts.get(selected_persona, system_prompts["Board / Executive Sponsor"])),
        ("human", prompt_template)
    ])

    # Extract vars manually (same as chain lambdas) to format the full prompt
    vars_for_prompt = {
        "tool_name": state["input_tool"].get("Tool Name", "Selected Tool"),
        "vendor": state["input_tool"].get("Vendor", "Unknown"),
        "persona": state.get("persona", "Board / Executive Sponsor"),
        "tone": PERSONA_PROMPTS[state.get("persona", "Board / Executive Sponsor")]["tone"],
        "focus": PERSONA_PROMPTS[state.get("persona", "Board / Executive Sponsor")]["focus"],
        "dominant_category": build_persona_focus_text(state.get("persona", "Board / Executive Sponsor"), weights)[0],
        "persona_focus_block": build_persona_focus_text(state.get("persona", "Board / Executive Sponsor"), weights)[1],
        "input_tool_data": json.dumps(state["input_tool"], separators=(',', ':')),
        "compare_tools_data": json.dumps(state["compare_tools"], separators=(',', ':'))
    }
    
    full_prompt = prompt.invoke(vars_for_prompt)  # This formats the template with vars
    prompt_str = "\n\n".join([f"{msg.type}: {msg.content}" for msg in full_prompt.messages]) 

    # js_code = f"console.log({json.dumps(prompt_str)});"
    # components.html(f"<script>{js_code}</script>", height=0)
     
    
    # print("=== FULL PROMPT DEBUG ===")
    # print(prompt_str)
    # print("=== END DEBUG ===")
    
    # with open("debug_prompt.txt", "w") as f:
    #     f.write(prompt_str)
    # print("Full prompt saved to debug_prompt.txt")


   

    chain = (
        {
            "tool_name": lambda x: x["input_tool"].get("Tool Name", "Selected Tool"),
            "vendor": lambda x: x["input_tool"].get("Vendor", "Unknown"),
            "persona": lambda x: x.get("persona", "Board / Executive Sponsor"),
            "tone": lambda x: PERSONA_PROMPTS[x.get("persona", "Board / Executive Sponsor")]["tone"],
            "focus": lambda x: PERSONA_PROMPTS[x.get("persona", "Board / Executive Sponsor")]["focus"],
            "dominant_category": lambda x: build_persona_focus_text(x.get("persona", "Board / Executive Sponsor"),x.get("weights"))[0],
            "persona_focus_block": lambda x: build_persona_focus_text(x.get("persona", "Board / Executive Sponsor"), x.get("weights"))[1], 
            "input_tool_data": lambda x: json.dumps(x["input_tool"], separators=(',', ':')),
            "compare_tools_data": lambda x: json.dumps(x["compare_tools"], separators=(',', ':'))
        }
        | prompt
        | chat_llm
        | StrOutputParser()
    )

    with get_openai_callback() as cb:
        response = chain.invoke(state)

    #rec_md = recommendations(st.session_state.selected_sheet)
    # response is the LLM-generated string; prepend recommendations +"\n\n"+"\n\n"+str(rec_md)
    final_report = str(response)

    state["actual_tokens"] = {
        "prompt_tokens": cb.prompt_tokens,
        "completion_tokens": cb.completion_tokens,
        "total_tokens": cb.total_tokens
    }
    state["report"] = final_report
    return state

# ==================== DOCX EXPORT (Improved Table Handling) =========================================================================
def create_docx_from_markdown(md_content):
    doc = Document()
    doc.add_heading('Research & Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line[2:-2], level=1)
        elif line.startswith("| **Strengths**") or line.startswith("| **Opportunities**"):
            # Start of SWOT table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            if len(table_lines) >= 6:
                doc.add_paragraph()
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'

                # Fill cells with preserved line breaks
                cells_content = [
                    [tl.split("|")[1].strip(), tl.split("|")[2].strip()] if len(tl.split("|")) > 2 else ["", ""]
                    for tl in table_lines[:6]
                ]

                for row_idx, row in enumerate(cells_content):
                    for col_idx, cell_text in enumerate(row):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ""
                        for bullet in cell_text.splitlines():
                            if bullet.strip():
                                p = cell.add_paragraph(bullet.strip(), style='List Bullet' if row_idx > 1 else None)
                                if row_idx <= 1:
                                    p.runs[0].bold = True
        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
        else:
            doc.add_paragraph(line)
        i += 1

    doc.add_page_break()
    footer = doc.add_paragraph(f'Generated on {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ==================== GRAPH & CACHING =================================================================================================
def build_graph():
    from langgraph.graph import StateGraph
    wf = StateGraph(ComparisonState)

    wf.add_node("prepare", prepare_data)
    wf.add_node("compare", generate_comparison)
    wf.set_entry_point("prepare")
    wf.add_edge("prepare", "compare")
    wf.add_edge("compare", END)
    return wf.compile()

graph = build_graph()


#@st.cache_data
def load_sheet_names(_excel_path):
    #Loading sheet names by taking excel path as input and returning list of sheet names
    xls = pd.ExcelFile(_excel_path)
    return xls.sheet_names

def load_tools(_excel_path, _sheet_name):
    #Loading tools by taking sheet name and excel path as input and returning dataframe
    df = pd.read_excel(_excel_path, sheet_name=_sheet_name)
    df.columns = df.columns.str.strip()
    return df

def generate_fresh_report(_input_tool_json, _compare_tools_json, _persona, _weights_json):
    """Generate report FRESH every time (no caching)"""
    input_tool = json.loads(_input_tool_json)
    compare_tools = json.loads(_compare_tools_json)
    weights = json.loads(_weights_json)
    
    initial_state = {
        "input_tool": input_tool,
        "compare_tools": compare_tools,
        "persona": _persona,
        "weights": weights,
        "weight_mode": "fresh",
        "report": "",
        "estimated_tokens": 0,
        "actual_tokens": {},
        "estimation_done": False
    }
    
    prep_state = prepare_data(initial_state)
    result = graph.invoke(prep_state)
    return result

def generate_fresh_comparison(_input_tool_json, _compare_tool_json, _persona, _weights_json):
    """Generate comparison report FRESH every time (no caching)"""
    input_tool = json.loads(_input_tool_json)
    alt_tool = json.loads(_compare_tool_json)
    weights = json.loads(_weights_json)
    
    new_state = {
        "input_tool": input_tool,
        "alternative_tool": alt_tool,
        "persona": _persona,
        "weights": weights,
        "weight_mode": "cached",
        "new_report": "",
        "estimated_tokens": 0,
        "actual_tokens": {},
        "estimation_done": False
    }
    
    prepare_state = preparation_data(new_state)
    new_result = graph1.invoke(prepare_state)
    return new_result


#kept for sake of reference, can be removed in final version

# # ==================== STREAMLIT UI =====================================================================================================
# st.set_page_config(layout="wide", page_title="RAR - Research & Analysis", page_icon="🔍")

# if "weight_mode" not in st.session_state:
#     st.session_state.weight_mode = "static(Default Weights)"

# if "active_weights" not in st.session_state:
#     st.session_state.active_weights = None  

# if "dynamic_weights" not in st.session_state:
#     st.session_state.dynamic_weights = None  

# if "dynamic_reasoning" not in st.session_state:
#     st.session_state.dynamic_reasoning = None

# if "prev_weight_mode" not in st.session_state:
#     st.session_state.prev_weight_mode = "static(Default Weights)"
                
# # Local report storage (not persisted in session - regenerated per run)
# _current_report = None
# _current_comparison_report = None

# if "comparison_mode" not in st.session_state:
#     st.session_state.comparison_mode = False

# if "compare_tool" not in st.session_state:
#     st.session_state.compare_tool = None

# if "report_generated" not in st.session_state:
#     st.session_state.report_generated = False

# if "comparison_report_generated" not in st.session_state:
#     st.session_state.comparison_report_generated = False

# if "last_report_cache_key" not in st.session_state:
#     st.session_state.last_report_cache_key = None

# if "comparison_result" not in st.session_state:
#     st.session_state.comparison_result = None

# if "main_report_result" not in st.session_state:
#     st.session_state.main_report_result = None

# st.title(" Research & Analysis Report Generator")

# for key in ['selected_sector', 'selected_sheet', 'input_choice', 'selected_persona', 'selected_category']:
#     if key not in st.session_state:
#         st.session_state[key] = None

# if st.session_state.selected_sector is None:
#     st.session_state.selected_sector = "Manufacturing"
# if st.session_state.selected_persona is None:
#     st.session_state.selected_persona = "Board / Executive Sponsor"


# col1, col2 = st.columns([1, 2])

# with col1:
#     st.subheader(" Configuration")

#     selected_sector = st.selectbox("Sector:", list(SECTOR_PATHS.keys()), 
#                                    index=list(SECTOR_PATHS.keys()).index(st.session_state.selected_sector))

#     if selected_sector != st.session_state.selected_sector:
#         st.session_state.selected_sector = selected_sector
#         st.session_state.selected_sheet = None
#         st.session_state.input_choice = None
#         st.session_state.report_generated = False
#         st.session_state.comparison_report_generated = False
#         st.session_state.main_report_result = None
#         st.rerun()

#     excel_path = SECTOR_PATHS.get(selected_sector)
    
#     if excel_path is None:
#         st.warning("Data in progress for this sector.")
#     else:
#         sheet_names = load_sheet_names(excel_path)
#         if st.session_state.selected_sheet is None:
#             st.session_state.selected_sheet = sheet_names[0]

#         selected_sheet = st.selectbox("Stage / Sheet:", sheet_names,
#                                       index=sheet_names.index(st.session_state.selected_sheet))

#         if selected_sheet != st.session_state.selected_sheet:
#             st.session_state.selected_sheet = selected_sheet
#             st.session_state.input_choice = None
#             st.session_state.report_generated = False
#             st.session_state.comparison_report_generated = False
#             st.session_state.main_report_result = None
#             st.rerun()

#         df = load_tools(excel_path, selected_sheet)
#         tools = df['Tool Name'].dropna().unique().tolist()

#         if len(tools) > 0:
#             if st.session_state.input_choice not in tools:
#                 st.session_state.input_choice = tools[0]

#             input_choice = st.selectbox("Tool to Analyze:", tools,
#                                         index=tools.index(st.session_state.input_choice),
#                                         key="tool_select")

#             if input_choice != st.session_state.input_choice:
#                 st.session_state.input_choice = input_choice
#                 st.session_state.report_generated = False
#                 st.session_state.comparison_report_generated = False
#                 st.session_state.main_report_result = None

#             st.markdown("---")
#             st.markdown("#### 👤 Report Audience")

#             persona_options = list(PERSONA_PROMPTS.keys())
#             selected_persona = st.selectbox(
#                 "Tailor report for:",
#                 options=persona_options,
#                 format_func=lambda x: f"{PERSONA_PROMPTS[x]['icon']} {x}",
#                 index=persona_options.index(st.session_state.selected_persona),
#                 help="Changes tone, depth, and focus of the entire report",
#                 key="persona_select_final"
#             )

#             if selected_persona != st.session_state.selected_persona:
#                 st.session_state.selected_persona = selected_persona
#                 st.session_state.report_generated = False
#                 st.session_state.comparison_report_generated = False
#                 st.session_state.main_report_result = None

#             category_options = list(persona_weights[selected_persona].keys())

#             if st.session_state.selected_category not in category_options:
#                 st.session_state.selected_category = category_options[0]
#             selected_category = st.selectbox(
#                 "Select Feature Category:", options=category_options,
#                 index=category_options.index(st.session_state.selected_category),
#                 key="category_select"
#              )

#             if selected_category != st.session_state.selected_category:
#                 st.session_state.selected_category = selected_category
#                 st.session_state.report_generated = False
#                 st.session_state.comparison_report_generated = False
#                 st.session_state.main_report_result = None

#             st.markdown("Weight Selection mode")

#             weight_mode = st.radio(
#                 "What weights you wants to select",
#                 options=["static(Default Weights)", "User Based", "Dynamic (AI Optimized)"],
#                 key="weight_mode",
#                 help="choose how category importance is calculated"
#             )
            

            
#             if weight_mode != st.session_state.prev_weight_mode:
#                 st.session_state.report_generated = False
#                 st.session_state.comparison_report_generated = False
#                 st.session_state.main_report_result = None
#                 st.session_state.prev_weight_mode = weight_mode
#                 st.toast(f"switched to '{weight_mode}' mode")

#                 if weight_mode != "Dynamic (AI Optimized)":
#                     st.session_state.dynamic_weights = None

#             st.markdown("Category weightages")
#             if st.session_state.get("weight_mode") == "User Based":
#                 def rebalance_weights(changed_category):
#                     active = st.session_state.active_weights
#                     prev= st.session_state.prev_weights

#                     new_value = st.session_state[f"weight_{selected_persona}_{changed_category}"]
#                     active[changed_category] = new_value

#                     remaining = 1.0 - new_value
#                     others = [c for c in active if c != changed_category]

#                     old_remaining = sum(prev[c] for c in others)
#                     if old_remaining == 0:
#                         return

#                     for c in others:
#                         active[c] = prev[c] * remaining / old_remaining
#                         st.session_state[f"weight_{selected_persona}_{c}"] = active[c]

#                     total = sum(active.values())
#                     for c in active:
#                         active[c] /= total
#                         st.session_state[f"weight_{selected_persona}_{c}"] = active[c]     

#                     st.session_state.prev_weights = active.copy()

#                 if "prev_persona_for_weights" not in st.session_state:
#                     st.session_state.prev_persona_for_weights = selected_persona


#                 if selected_persona != st.session_state.prev_persona_for_weights:
#                     for k in list(st.session_state.keys()):
#                         if k.startswith("weight_") and k != "weight_mode":
#                             del st.session_state[k]
                    
#                     st.session_state.active_weights = persona_weights[selected_persona].copy()
#                     st.session_state.prev_weights = persona_weights[selected_persona].copy() 
#                     st.session_state.prev_persona_for_weights = selected_persona
#                     st.session_state.report_generated = False
#                     st.session_state.comparison_report_generated = False
#                     st.session_state.main_report_result = None
                    
#                 if st.session_state.get("weight_mode") == "User Based":
#                     if "active_weights" not in st.session_state or st.session_state.active_weights is None:
#                         st.session_state.active_weights = persona_weights[selected_persona].copy()
#                     if "prev_weights" not in st.session_state or st.session_state.prev_weights is None:
#                         st.session_state.prev_weights = st.session_state.active_weights.copy()

                    
#                 changed_category = None
#                 for category, default_weight in st.session_state.active_weights.items():
#                     value = st.number_input(
#                         label=category,
#                         min_value=0.0,
#                         max_value=1.0,
#                         value=st.session_state.active_weights[category],
#                         step=0.01,
#                         key=f"weight_{selected_persona}_{category}",
#                         on_change=rebalance_weights,
#                         args=(category,)
#                     )

#             effective_weights = None
            
#             tool_row = df[df['Tool Name'] == input_choice].iloc[0]
#             st.markdown(f"**Selected Tool:** {tool_row['Tool Name']} ({tool_row['Vendor']})")
            
#             # Define columns once for reuse
#             cols = ['Tool Name', 'Tool Type', 'Vendor', 'Region', 'Cost', 'Number of Integrations', 'Efficiency', 'Version', 'Core Features', 'Customizability', 'Integration Capability', 'Multi-Modal Support', 'Automation', 'Knowledge Management', 'Speed / Latency', 'Scalability', 'Robustness', 'Adaptability', 'Energy Efficiency', 'UX/UI', 'Accessibility', 
# 'Support & Community', 'Gamification', 'Architecture', 'Security & Privacy', 'Interoperability', 'Updates & Maintenance',
# 'Resilience Engineering', 'Geo-Sovereignty', 'Licensing Model', 'TCO (Total Cost of Ownership)', 'ROI / Value Creation', 'Vendor Stability', 'Procurement Flexibility', 'Cash Flow Impact',
# 'Revenue Enablement', 'Cost Optimization', 'Profitability Impact', 'Risk Exposure', 'Sustainable Finance Impact', 'Competitive Advantage', 'Market Fit', 'Time-to-Market', 
# 'Ecosystem Partnerships', 'Innovation Potential', 'Customer Retention Impact', 'Global Expansion Enabler', 'Brand Reputation Impact', 'Data Governance', 'Ethical AI Practices', 'Cyber Resilience', 'Business Continuity', 'Geopolitical Risk', 'Collaboration Features', 'Customer/User Value', 'Futuristic Readiness']

#             compare_button = st.button(" Generate Report", type="primary", use_container_width=True)

#             if compare_button:
#                 #
#                 with st.spinner(f"Generating report for {selected_persona}..."):
#                     if st.session_state.weight_mode == "Dynamic (AI Optimized)":
#                         if st.session_state.dynamic_weights is None:
#                             final_weights, reasoning = run_ai_weight_optimization(
#                                 sector=selected_sector,
#                                 stage=selected_sheet,
#                                 tool=input_choice,
#                                 persona=selected_persona,
#                                 feature_categories=selected_category,
#                                 initial_weights=persona_weights[selected_persona]
#                                 )
#                             st.session_state.dynamic_weights = final_weights
#                             st.session_state.dynamic_reasoning = reasoning
#                         effective_weights = st.session_state.dynamic_weights
#                     else:
#                         effective_weights = st.session_state.active_weights  
                    
#                     weight_mode = st.session_state.get("weight_mode", "static(Default Weights)")
#                     active_weights = st.session_state.get("active_weights", None)
#                     dynamic_weights = st.session_state.get("dynamic_weights", None)

#                     effective_weights = get_effective_weights(
#                         selected_persona=selected_persona,
#                         weight_mode=weight_mode,
#                         active_weights=active_weights,
#                         dynamic_weights=dynamic_weights
#                     )

#                     input_tool = tool_row[cols].to_dict()
#                     compare_tools = df[df['Tool Name'] != input_choice][cols].to_dict('records')

#                     # Generate fresh main report ALWAYS
#                     result = generate_fresh_report(
#                         json.dumps(input_tool),
#                         json.dumps(compare_tools),
#                         selected_persona,
#                         json.dumps(effective_weights)
#                     )
#                     # Store result in session state
#                     st.session_state.main_report_result = result
#                     st.session_state.report_generated = True
                    
#                     st.rerun()

#             st.markdown("---")
#             st.subheader(" Comparison Between Tools")
#             all_tools = df['Tool Name'].dropna().unique().tolist()

#             compare_options = [t for t in all_tools if t != st.session_state.input_choice]
#             if compare_options:
#                 if st.session_state.compare_tool not in compare_options:
#                     st.session_state.compare_tool = compare_options[0]

#                 selected_compare_tool = st.selectbox(
#                     "Select tool to compare with:",
#                     options=compare_options,
#                     index=compare_options.index(st.session_state.compare_tool),
#                     key="compare_tool_select"
#                 )

#                 if selected_compare_tool != st.session_state.compare_tool:
#                     st.session_state.compare_tool = selected_compare_tool
#                     st.session_state.comparison_report_generated = False
#                     st.session_state.last_comparison_cache_key = None
#                 else:
#                     st.session_state.compare_tool = selected_compare_tool

#                 compare_btn = st.button(
#                     f"compare {st.session_state.input_choice} vs {st.session_state.compare_tool}",
#                     type="secondary",
#                     use_container_width=True
#                 )

#                 if compare_btn:
#                     with st.spinner(f"Comparison between tools: {st.session_state.input_choice} vs {st.session_state.compare_tool}"):
#                         tool_row = df[df["Tool Name"] == st.session_state.input_choice].iloc[0]
#                         alt_row = df[df['Tool Name'] == st.session_state.compare_tool].iloc[0]

#                         weight_mode = st.session_state.get("weight_mode", "static(Default Weights)")
#                         active_weights = st.session_state.get("active_weights", None)
#                         dynamic_weights = st.session_state.get("dynamic_weights", None)

#                         effective_weights = get_effective_weights(
#                             selected_persona=selected_persona,
#                             weight_mode=weight_mode,
#                             active_weights=active_weights,
#                             dynamic_weights=dynamic_weights
#                         )

#                         input_tool = tool_row[cols].to_dict()
#                         alternative_tool = alt_row[cols].to_dict()

#                         # Generate fresh comparison report ALWAYS
#                         new_result = generate_fresh_comparison(
#                             json.dumps(input_tool),
#                             json.dumps(alternative_tool),
#                             selected_persona,
#                             json.dumps(effective_weights)
#                         )
#                         # Store result in session state
#                         st.session_state.comparison_result = new_result
#                         st.session_state.comparison_mode = True
#                         st.session_state.comparison_report_generated = True
                        
#                         st.rerun()
                

# with col2:
#     st.subheader(" Generated Report")

#     if st.session_state.report_generated and hasattr(st.session_state, 'main_report_result') and st.session_state.main_report_result:
#         # Use freshly generated result from button click
#         result = st.session_state.main_report_result
#         report = result["report"]

#         st.success(f"Report generated for {st.session_state.selected_persona}!")

#         docx_bytes = create_docx_from_markdown(report)
#         st.download_button(
#             "📥 Download as Word (.docx)",
#             data=docx_bytes,
#             file_name=f"RAR_{st.session_state.input_choice.replace(' ', '_')}_{st.session_state.selected_persona.replace(' ', '_').replace('/', '')}.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             use_container_width=True
#         )

#         st.markdown("---")
#         with st.container(border=True):
#             st.markdown(report, unsafe_allow_html=True)
#     else:
#         st.info("Select a tool and persona, then click **Generate Report** to see the output!")

#     if st.session_state.comparison_mode:
#         st.markdown("---")

#         if st.session_state.comparison_report_generated and hasattr(st.session_state, 'comparison_result') and st.session_state.comparison_result:
#             # Use freshly generated result from button click
#             new_result = st.session_state.comparison_result
#             new_report = new_result["new_report"]
            
#             st.success(f"Comparison Report generated for {st.session_state.input_choice} vs {st.session_state.compare_tool}!")
            
#             st.info(" **Scroll down below to view the comparison report**")

#             docx_bytes = create_docx_from_markdown(new_report)
#             st.download_button(
#                 "📥 Download Comparison Report (.docx)",
#                 data=docx_bytes,
#                 file_name=f"VS_{st.session_state.input_choice}_vs_{st.session_state.compare_tool}.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 use_container_width=True
#                 )
#             st.markdown("---")
#             with st.container(border=True):
#                 st.markdown(new_report, unsafe_allow_html=True)
#     else:
#         st.info("Select selected tool and alternative tool, then click **Compare_btn** to see the output!")    

# if st.session_state.comparison_mode and st.session_state.comparison_report_generated and hasattr(st.session_state, 'comparison_result') and st.session_state.comparison_result:
#     st.sidebar.info(" **Comparison Report Ready!** \n\nScroll down in the main panel to see the detailed comparison report between:\n• " + st.session_state.input_choice + "\n• " + st.session_state.compare_tool)
# else:
#     st.sidebar.success("Persona-tailored structured reports ready!")